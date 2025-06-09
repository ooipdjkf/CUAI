from docx import Document
from pathlib import Path
from typing import Dict, List, Tuple
import re
from collections import Counter

class DocumentAnalyzer:
    def __init__(self):
        self.section_patterns = []
        self.formula_patterns = []
        self.table_count = 0
        self.section_levels = Counter()
        self.special_chars = Counter()
        
    def analyze_document(self, doc_path: str) -> Dict:
        """分析Word文档的结构"""
        print(f"\n开始分析文档: {Path(doc_path).name}")
        doc = Document(doc_path)
        analysis = {
            'filename': Path(doc_path).name,
            'total_paragraphs': len(doc.paragraphs),
            'total_tables': len(doc.tables),
            'sections': self._analyze_sections(doc),
            'formulas': self._analyze_formulas(doc),
            'tables': self._analyze_tables(doc),
            'special_patterns': self._find_special_patterns(doc)
        }
        return analysis
        
    def _analyze_sections(self, doc: Document) -> Dict:
        """分析文档的章节结构"""
        sections = []
        current_section = None
        
        for para in doc.paragraphs:
            # 分析段落样式和编号
            style_name = para.style.name if para.style else 'Normal'
            text = para.text.strip()
            
            if text:
                # 尝试识别章节编号模式
                section_match = re.match(r'^(\d+\.)*\d+\s+', text)
                if section_match:
                    section_number = section_match.group(0).strip()
                    section_title = text[len(section_number):].strip()
                    sections.append({
                        'number': section_number,
                        'title': section_title,
                        'level': section_number.count('.') + 1,
                        'style': style_name
                    })
                    
        return {
            'total_sections': len(sections),
            'section_levels': Counter(s['level'] for s in sections),
            'section_samples': sections[:5]  # 前5个章节作为样本
        }
        
    def _analyze_formulas(self, doc: Document) -> Dict:
        """分析文档中的公式"""
        formulas = []
        formula_patterns = []
        
        for para in doc.paragraphs:
            text = para.text.strip()
            # 查找可能的公式模式
            formula_matches = re.finditer(r'\(([\d.]+)\)(.*?)(?=\(|$)', text)
            for match in formula_matches:
                formula_id = match.group(1)
                expression = match.group(2).strip()
                formulas.append({
                    'id': formula_id,
                    'expression': expression,
                    'context': text
                })
                
        return {
            'total_formulas': len(formulas),
            'formula_samples': formulas[:5],  # 前5个公式作为样本
            'special_characters': self._extract_special_chars(formulas)
        }
        
    def _analyze_tables(self, doc: Document) -> Dict:
        """分析文档中的表格"""
        tables_info = []
        
        for i, table in enumerate(doc.tables):
            if not table.rows:
                continue
                
            # 获取表格的基本信息
            table_info = {
                'index': i + 1,
                'rows': len(table.rows),
                'columns': len(table.columns),
                'sample_content': []
            }
            
            # 获取表格内容样本（最多3行）
            for row_idx in range(min(3, len(table.rows))):
                row_content = []
                for cell in table.rows[row_idx].cells:
                    cell_text = cell.text.strip()
                    row_content.append(cell_text)
                table_info['sample_content'].append(row_content)
            
            tables_info.append(table_info)
            
        return {
            'total_tables': len(tables_info),
            'table_samples': tables_info[:3]  # 前3个表格作为样本
        }
        
    def _find_special_patterns(self, doc: Document) -> Dict:
        """查找文档中的特殊模式"""
        patterns = {
            'references': [],
            'units': set(),
            'special_symbols': set()
        }
        
        for para in doc.paragraphs:
            text = para.text
            # 查找引用
            ref_matches = re.finditer(r'(?:según|véase|ver)\s+(?:el\s+)?(?:apartado|punto)\s+([\d.]+)', text, re.IGNORECASE)
            patterns['references'].extend(m.group(1) for m in ref_matches)
            
            # 查找单位
            unit_matches = re.finditer(r'\b\d+(?:,\d+)?\s*(mm|cm|m|kN|MPa|°C|N/mm²)\b', text)
            patterns['units'].update(m.group(1) for m in unit_matches)
            
            # 查找特殊符号
            special_chars = re.findall(r'[^\w\s,.;:()\[\]{}]', text)
            patterns['special_symbols'].update(special_chars)
            
        return patterns
        
    def _extract_special_chars(self, formulas: List[Dict]) -> List[str]:
        """提取公式中的特殊字符"""
        special_chars = set()
        for formula in formulas:
            chars = re.findall(r'[^\w\s,.;:()\[\]{}]', formula['expression'])
            special_chars.update(chars)
        return list(special_chars)

def main():
    """主函数"""
    analyzer = DocumentAnalyzer()
    
    # 分析两个文档
    docs = [
        r"H:\trabajo\recien works\OPEN structure\Code consultant\DBSE-A.docx",
        r"H:\trabajo\recien works\OPEN structure\Code consultant\RSCIEI.docx"
    ]
    
    for doc_path in docs:
        try:
            analysis = analyzer.analyze_document(doc_path)
            
            # 打印分析结果
            print("\n文档基本信息:")
            print(f"总段落数: {analysis['total_paragraphs']}")
            print(f"总表格数: {analysis['total_tables']}")
            
            print("\n章节结构:")
            print(f"总章节数: {analysis['sections']['total_sections']}")
            print("章节层级分布:", dict(analysis['sections']['section_levels']))
            
            if analysis['sections']['section_samples']:
                print("\n章节样本:")
                for section in analysis['sections']['section_samples'][:3]:
                    print(f"  {section['number']} {section['title']}")
            
            print("\n公式分析:")
            print(f"总公式数: {analysis['formulas']['total_formulas']}")
            if analysis['formulas']['formula_samples']:
                print("公式样本:")
                for f in analysis['formulas']['formula_samples'][:2]:
                    print(f"  ID: {f['id']}, 表达式: {f['expression']}")
            
            if analysis['tables']['table_samples']:
                print("\n表格样本:")
                for table in analysis['tables']['table_samples'][:1]:
                    print(f"  表格 {table['index']} ({table['rows']}行 x {table['columns']}列)")
                    for row in table['sample_content'][:2]:
                        print(f"    {' | '.join(cell[:50] for cell in row)}")
                    
            print("\n特殊模式:")
            print("引用示例:", analysis['special_patterns']['references'][:5])
            print("单位:", list(analysis['special_patterns']['units']))
            print("特殊符号:", list(analysis['special_patterns']['special_symbols'])[:10])
            
        except Exception as e:
            print(f"处理文档 {doc_path} 时出错: {str(e)}")
            import traceback
            print(traceback.format_exc())

if __name__ == "__main__":
    main() 