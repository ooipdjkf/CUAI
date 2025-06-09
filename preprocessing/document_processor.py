from docx import Document as DocxDocument
from pathlib import Path
import re
from typing import Dict, List, Tuple, Optional
from collections import Counter
from preprocessing.data_structures import Document, Section, Formula, Table, Requirement, RequirementType
import json
import pickle
from datetime import datetime

class DocumentProcessor:
    def __init__(self):
        self.current_document: Optional[Document] = None
        self.current_section: Optional[Section] = None
        
    def process_document(self, doc_path: str) -> Document:
        """处理Word文档"""
        print(f"\n开始处理文档: {Path(doc_path).name}")
        docx = DocxDocument(doc_path)
        
        # 创建文档对象
        self.current_document = Document(filename=Path(doc_path).name)
        
        # 处理文档内容
        self._process_sections(docx)
        self._process_formulas(docx)
        self._process_tables(docx)
        self._process_requirements(docx)
        self._build_references()
        
        return self.current_document
        
    def _process_sections(self, docx: DocxDocument):
        """处理文档章节"""
        current_sections: Dict[int, Section] = {}  # level -> Section
        section_content = []  # 临时存储章节内容
        
        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 识别章节标题
            # 改进的章节匹配模式
            section_match = re.match(r'^(?P<number>(?:\d+\.)*\d+)\s+(?P<title>[A-Z].*$)', text)
            if section_match:
                # 如果有累积的内容，添加到当前章节
                if section_content and self.current_section:
                    self.current_section.content_es = "\n".join(section_content)
                section_content = []
                
                section_number = section_match.group('number').strip()
                section_title = section_match.group('title').strip()
                level = section_number.count('.') + 1
                
                # 创建新章节
                section = Section(
                    id=section_number,
                    title_es=section_title,
                    level=level
                )
                
                # 设置父子关系
                if level > 1:
                    # 查找父章节
                    parent_number = '.'.join(section_number.split('.')[:-1])
                    if parent_number in self.current_document.sections:
                        parent = self.current_document.sections[parent_number]
                        section.parent_section = parent.id
                        parent.subsections.append(section.id)
                
                # 更新当前章节
                self.current_document.sections[section.id] = section
                self.current_section = section
                current_sections[level] = section
            else:
                # 累积章节内容
                section_content.append(text)
        
        # 处理最后一个章节的内容
        if section_content and self.current_section:
            self.current_section.content_es = "\n".join(section_content)
            
    def _process_formulas(self, docx: DocxDocument):
        """处理公式"""
        current_context = []  # 存储公式上下文
        
        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 更新上下文
            current_context.append(text)
            if len(current_context) > 3:  # 保持3段上下文
                current_context.pop(0)
            
            # 查找公式
            formula_matches = re.finditer(r'\(([\d.]+)\)(.*?)(?=\(|$)', text)
            for match in formula_matches:
                formula_id = match.group(1)
                expression = match.group(2).strip()
                
                # 获取完整上下文
                context = "\n".join(current_context)
                
                # 创建公式对象
                formula = Formula(
                    id=formula_id,
                    expression=expression,
                    context_es=context,
                    section_id=self.current_section.id if self.current_section else ""
                )
                
                # 提取变量和单位
                self._extract_formula_details(formula)
                
                # 添加到文档和当前章节
                self.current_document.formulas[formula.id] = formula
                if self.current_section:
                    self.current_section.formulas.append(formula)
                    
    def _extract_formula_details(self, formula: Formula):
        """提取公式的变量和单位"""
        # 改进变量说明匹配
        var_patterns = [
            r'donde\s+(\w+)\s+es\s+([^,\.]+)',  # donde x es ...
            r'siendo\s+(\w+)\s+(?:el|la)\s+([^,\.]+)',  # siendo x el/la ...
            r'con\s+(\w+)\s+(?:el|la)\s+([^,\.]+)'  # con x el/la ...
        ]
        
        for pattern in var_patterns:
            var_matches = re.finditer(pattern, formula.context_es, re.IGNORECASE)
            for match in var_matches:
                var_name = match.group(1)
                var_desc = match.group(2).strip()
                formula.variables[var_name] = var_desc
        
        # 改进单位匹配
        unit_patterns = [
            r'\b\d*[,.]?\d*\s*(mm|cm|m|kN|MPa|°C|N/mm²|kN/m|kN·m|m²|m³)\b',
            r'\[([^]]+)\]'  # 匹配方括号中的单位
        ]
        
        units = set()
        for pattern in unit_patterns:
            unit_matches = re.finditer(pattern, formula.context_es)
            units.update(m.group(1) for m in unit_matches)
        
        formula.units = list(units)
        
    def _process_tables(self, docx: DocxDocument):
        """处理表格"""
        for i, table in enumerate(docx.tables, 1):
            if not table.rows:
                continue
                
            # 尝试获取表格标题
            caption = self._find_table_caption(table, i)
            
            # 创建表格对象
            table_obj = Table(
                id=f"Table_{i}",
                caption_es=caption,
                section_id=self.current_section.id if self.current_section else "",
                headers=self._extract_headers(table),
                data=self._extract_table_data(table)
            )
            
            # 添加到文档和当前章节
            self.current_document.tables[table_obj.id] = table_obj
            if self.current_section:
                self.current_section.tables.append(table_obj)
                
    def _find_table_caption(self, table, index: int) -> str:
        """查找表格标题"""
        # 实现表格标题查找逻辑
        return f"Tabla {index}"
        
    def _extract_headers(self, table) -> List[str]:
        """提取表格表头"""
        if not table.rows:
            return []
            
        headers = []
        for cell in table.rows[0].cells:
            headers.append(cell.text.strip())
        return headers
        
    def _extract_table_data(self, table) -> List[List[str]]:
        """提取表格数据"""
        data = []
        for row in table.rows[1:]:  # 跳过表头行
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.strip())
            data.append(row_data)
        return data
        
    def _process_requirements(self, docx: DocxDocument):
        """处理规范要求"""
        requirement_indicators = {
            RequirementType.MANDATORY: [
                r'debe(?:rá)?[n]?\s',
                r'será[n]?\s',
                r'es necesario\s',
                r'se requiere\s',
                r'ha(?:n)?\sde\s',
                r'no se admite[n]?\s',
                r'es obligatorio\s'
            ],
            RequirementType.RECOMMENDED: [
                r'se recomienda\s',
                r'es recomendable\s',
                r'conviene\s',
                r'es aconsejable\s',
                r'preferentemente\s'
            ],
            RequirementType.CONDITIONAL: [
                r'si\s.*?,\s(?:debe|deberá|será|es necesario)',
                r'cuando\s.*?,\s(?:debe|deberá|será|es necesario)',
                r'en(?:\sel)?\scaso\sde\s.*?,\s(?:debe|deberá|será|es necesario)',
                r'siempre\sque\s.*?,\s(?:debe|deberá|será|es necesario)'
            ]
        }
        
        current_context = []
        
        for para in docx.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 更新上下文
            current_context.append(text)
            if len(current_context) > 3:
                current_context.pop(0)
            
            # 检查是否包含要求
            for req_type, patterns in requirement_indicators.items():
                for pattern in patterns:
                    if re.search(pattern, text, re.IGNORECASE):
                        # 创建要求对象
                        requirement = Requirement(
                            text_es=text,
                            type=req_type
                        )
                        
                        # 设置上下文
                        requirement.context = "\n".join(current_context)
                        
                        # 查找相关公式和表格
                        self._find_related_items(requirement)
                        
                        # 提取条件（对于条件性要求）
                        if req_type == RequirementType.CONDITIONAL:
                            self._extract_conditions(requirement)
                        
                        # 添加到当前章节
                        if self.current_section:
                            self.current_section.requirements.append(requirement)
                        break
                        
    def _extract_conditions(self, requirement: Requirement):
        """提取要求的条件"""
        condition_patterns = [
            (r'si\s(.*?),\s', '如果'),
            (r'cuando\s(.*?),\s', '当'),
            (r'en(?:\sel)?\scaso\sde\s(.*?),\s', '在...的情况下'),
            (r'siempre\sque\s(.*?),\s', '只要')
        ]
        
        for pattern, prefix in condition_patterns:
            match = re.search(pattern, requirement.text_es, re.IGNORECASE)
            if match:
                condition = match.group(1).strip()
                requirement.conditions.append(f"{prefix}{condition}")
                
    def _find_related_items(self, requirement: Requirement):
        """查找要求相关的公式和表格"""
        # 改进公式引用匹配
        formula_patterns = [
            r'(?:ecuación|fórmula|expresión)\s*\(?([\d.]+)\)?',
            r'según\s+(?:la\s+)?(?:ecuación|fórmula|expresión)\s*\(?([\d.]+)\)?'
        ]
        
        formula_refs = set()
        for pattern in formula_patterns:
            matches = re.finditer(pattern, requirement.text_es, re.IGNORECASE)
            formula_refs.update(m.group(1) for m in matches)
        requirement.related_formulas = list(formula_refs)
        
        # 改进表格引用匹配
        table_patterns = [
            r'(?:tabla|cuadro)\s*\(?([\d.]+)\)?',
            r'según\s+(?:la\s+)?(?:tabla|cuadro)\s*\(?([\d.]+)\)?'
        ]
        
        table_refs = set()
        for pattern in table_patterns:
            matches = re.finditer(pattern, requirement.text_es, re.IGNORECASE)
            table_refs.update(m.group(1) for m in matches)
        requirement.related_tables = list(table_refs)
        
    def _build_references(self):
        """构建交叉引用关系"""
        # 处理章节间的引用
        for section in self.current_document.sections.values():
            # 查找文本中的引用
            ref_matches = re.finditer(
                r'(?:según|véase|ver)\s+(?:el\s+)?(?:apartado|punto)\s+([\d.]+)',
                section.content_es,
                re.IGNORECASE
            )
            
            for match in ref_matches:
                ref_id = match.group(1)
                if ref_id in self.current_document.sections:
                    # 添加引用关系
                    section.references.append(ref_id)
                    self.current_document.sections[ref_id].referenced_by.append(section.id)
                    
        # 处理公式引用
        for formula in self.current_document.formulas.values():
            if formula.section_id:
                section = self.current_document.sections.get(formula.section_id)
                if section:
                    # 查找其他章节对该公式的引用
                    for other_section in self.current_document.sections.values():
                        if re.search(
                            fr'(?:ecuación|fórmula|expresión)\s*\(?{formula.id}\)?',
                            other_section.content_es,
                            re.IGNORECASE
                        ):
                            formula.referenced_by.append(other_section.id)

class DocumentEncoder(json.JSONEncoder):
    """用于JSON序列化文档对象"""
    def default(self, obj):
        if isinstance(obj, set):
            return list(obj)
        if hasattr(obj, '__dict__'):
            return {
                '_type': obj.__class__.__name__,
                'data': {k: v for k, v in obj.__dict__.items() if not k.startswith('_')}
            }
        return super().default(obj)

def save_processed_data(document: Document, output_dir: str):
    """保存处理后的数据"""
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)
    
    # 生成时间戳
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = Path(document.filename).stem
    
    # 保存JSON格式
    json_path = output_path / f"{base_name}_{timestamp}.json"
    with open(json_path, 'w', encoding='utf-8') as f:
        json.dump(document, f, cls=DocumentEncoder, ensure_ascii=False, indent=2)
    
    # 保存pickle格式（保留完整的对象结构）
    pickle_path = output_path / f"{base_name}_{timestamp}.pkl"
    with open(pickle_path, 'wb') as f:
        pickle.dump(document, f)
    
    print(f"\n数据已保存:")
    print(f"JSON文件: {json_path}")
    print(f"Pickle文件: {pickle_path}")
    
    # 生成统计信息
    stats = {
        'filename': document.filename,
        'total_sections': len(document.sections),
        'total_formulas': len(document.formulas),
        'total_tables': len(document.tables),
        'section_levels': {},
        'requirement_types': {
            'mandatory': 0,
            'recommended': 0,
            'conditional': 0
        }
    }
    
    # 统计章节级别
    for section in document.sections.values():
        stats['section_levels'][section.level] = stats['section_levels'].get(section.level, 0) + 1
        
        # 统计要求类型
        for req in section.requirements:
            stats['requirement_types'][req.type.value] += 1
    
    # 保存统计信息
    stats_path = output_path / f"{base_name}_{timestamp}_stats.json"
    with open(stats_path, 'w', encoding='utf-8') as f:
        json.dump(stats, f, ensure_ascii=False, indent=2)
    
    print(f"统计信息: {stats_path}")
    
    return json_path, pickle_path, stats_path

def main():
    """主函数"""
    processor = DocumentProcessor()
    
    # 处理两个文档
    docs = [
        "docs/DBSE-A.docx",
        "docs/RSCIEI.docx"
    ]
    
    output_dir = "processed_data"
    
    for doc_path in docs:
        try:
            document = processor.process_document(doc_path)
            
            # 打印处理结果
            print(f"\n文档: {document.filename}")
            print(f"总章节数: {len(document.sections)}")
            print(f"总公式数: {len(document.formulas)}")
            print(f"总表格数: {len(document.tables)}")
            
            # 打印一些示例章节
            print("\n章节示例:")
            for section_id, section in list(document.sections.items())[:3]:
                print(f"\n章节 {section.id}: {section.title_es}")
                print(f"  级别: {section.level}")
                print(f"  子章节数: {len(section.subsections)}")
                print(f"  公式数: {len(section.formulas)}")
                print(f"  表格数: {len(section.tables)}")
                print(f"  要求数: {len(section.requirements)}")
                
            # 打印一些示例公式
            print("\n公式示例:")
            for formula_id, formula in list(document.formulas.items())[:2]:
                print(f"\n公式 {formula.id}:")
                print(f"  表达式: {formula.expression}")
                if formula.variables:
                    print(f"  变量: {formula.variables}")
                if formula.units:
                    print(f"  单位: {formula.units}")
                    
            # 打印一些示例要求
            requirements_found = False
            for section in document.sections.values():
                if section.requirements:
                    print("\n要求示例:")
                    for req in section.requirements[:2]:
                        print(f"\n类型: {req.type.value}")
                        print(f"内容: {req.text_es[:100]}...")
                        if req.related_formulas:
                            print(f"相关公式: {req.related_formulas}")
                        if req.related_tables:
                            print(f"相关表格: {req.related_tables}")
                    requirements_found = True
                    break
                    
            if not requirements_found:
                print("\n未找到要求示例")
            
            # 保存处理后的数据
            save_processed_data(document, output_dir)
            
        except Exception as e:
            print(f"处理文档 {doc_path} 时出错: {str(e)}")
            import traceback
            print(traceback.format_exc())

if __name__ == "__main__":
    main() 